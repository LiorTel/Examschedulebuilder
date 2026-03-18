from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from phase_a.hebrew_utils import detect_semester_from_text, normalize_hebrew_text


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv"}
DATE_HINT_PATTERN = re.compile(r"\d{1,2}[./-]\d{1,2}(?:[./-]\d{2,4})?|(\d{1,2}\s*,\s*)+\d{1,2}\s+ב?[א-ת]+\s+\d{4}|\b(?:ב)?(?:ינואר|פברואר|מרץ|מרס|אפריל|מאי|יוני|יולי|אוגוסט|ספטמבר|אוקטובר|נובמבר|דצמבר)\b")
EVENT_START_KEYWORDS = [
    "היום הראשון",
    "היום האחרון",
    "תקופת בחינות",
    "תחילת תקופת הבחינות",
    "ימי היערכות",
    "חופשת",
    "פגרת",
    "יום הסטודנט",
    "יום פתוח",
    "יום אוריינטציה",
    "יום השלמה",
    "הפסקת לימודים",
    "הלימודים יסתיימו",
    "בחינה פסיכומטרית",
    'בחינות "סמסטר ראשון בתיכון"',
    "פתיחת שנה",
]
SECTION_KEYWORDS = ["אין לשבץ בחינות", "חגים ומועדים נוספים"]


@dataclass
class ParsedDocument:
    raw_text: str
    table_rows: list[str]
    blocks: list[dict]


def _table_rows_from_dataframe(df) -> Iterable[str]:
    for _, row in df.fillna("").iterrows():
        values = [str(v).strip() for v in row.tolist() if str(v).strip()]
        if values:
            yield " | ".join(values)


def _is_semester_heading(line: str) -> bool:
    clean = normalize_hebrew_text(line)
    return clean.startswith("סמסטר") and len(clean.split()) <= 3 and not DATE_HINT_PATTERN.search(clean)


def _is_section_heading(line: str) -> bool:
    clean = normalize_hebrew_text(line)
    return any(k in clean for k in SECTION_KEYWORDS) and clean.endswith(":")


def _is_event_start(line: str) -> bool:
    if any(keyword in line for keyword in EVENT_START_KEYWORDS):
        return True
    if "סמסטר" in line and any(k in line for k in ["היום הראשון", "היום האחרון", "תקופת", "תחילת"]):
        return True
    return False


def _build_blocks(lines: list[str]) -> list[dict]:
    blocks: list[dict] = []
    current_semester = "NONE"
    current_section = "GENERAL"
    in_no_exam_section = False
    current: dict | None = None

    def flush_current() -> None:
        nonlocal current
        if not current:
            return
        current["text"] = "\n".join(current["lines"])
        blocks.append(current)
        current = None

    for raw_line in lines:
        line = normalize_hebrew_text(raw_line)
        if not line:
            continue

        if _is_semester_heading(line):
            current_semester = detect_semester_from_text(line) or "NONE"
            in_no_exam_section = False
            current_section = "GENERAL"
            flush_current()
            continue

        if _is_section_heading(line):
            flush_current()
            current_section = line.rstrip(":")
            in_no_exam_section = "אין לשבץ בחינות" in line
            continue

        if _is_event_start(line):
            flush_current()
            current = {
                "title": line,
                "lines": [line],
                "semester_context": detect_semester_from_text(line) or current_semester,
                "in_no_exam_section": in_no_exam_section,
                "section_context": current_section,
            }
            continue

        if current is None:
            current = {
                "title": line,
                "lines": [line],
                "semester_context": detect_semester_from_text(line) or current_semester,
                "in_no_exam_section": in_no_exam_section,
                "section_context": current_section,
            }
        else:
            current["lines"].append(line)

    flush_current()
    return blocks


def parse_uploaded_file(file_name: str, file_bytes: bytes) -> ParsedDocument:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")
    if not file_bytes:
        raise ValueError("Uploaded file is empty")

    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        raw_text = "\n".join(pages)
        table_rows: list[str] = []
    elif extension == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        raw_text = "\n".join(p.text for p in doc.paragraphs)
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    table_rows.append(" | ".join(values))
    elif extension == ".xlsx":
        import pandas as pd

        workbook = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        table_rows = []
        text_blocks = []
        for sheet_name, df in workbook.items():
            text_blocks.append(f"Sheet: {sheet_name}")
            parsed_rows = list(_table_rows_from_dataframe(df))
            text_blocks.extend(parsed_rows)
            table_rows.extend(parsed_rows)
        raw_text = "\n".join(text_blocks)
    else:  # .csv
        import pandas as pd

        df = pd.read_csv(io.BytesIO(file_bytes))
        table_rows = list(_table_rows_from_dataframe(df))
        raw_text = "\n".join(table_rows)

    lines = [ln for ln in raw_text.splitlines() if ln.strip()]
    blocks = _build_blocks(lines)
    return ParsedDocument(raw_text=raw_text, table_rows=table_rows, blocks=blocks)
